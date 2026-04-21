#!/usr/bin/env python3
"""
gerar_docx.py — Gerador de currículo profissional em .docx
Uso: python scripts/gerar_docx.py --dados dados.json --saida curriculo.docx [--estilo corporativo]

Estilos disponíveis: minimalista | corporativo | criativo | executivo
"""

import argparse
import json
import sys
from pathlib import Path

try:
    from docx import Document
    from docx.shared import Pt, Cm, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
except ImportError:
    print("❌ Biblioteca python-docx não encontrada.")
    print("   Instale com: pip install python-docx")
    sys.exit(1)


# ─── PALETAS DE CORES POR ESTILO ───────────────────────────────────────────
ESTILOS = {
    "corporativo": {
        "primaria": RGBColor(26, 43, 74),     # Azul escuro profissional
        "secundaria": RGBColor(52, 90, 145),   # Azul médio
        "texto": RGBColor(30, 30, 30),         # Quase preto
        "sutil": RGBColor(100, 110, 130),      # Cinza azulado
        "fonte_titulo": "Calibri Light",
        "fonte_corpo": "Calibri",
    },
    "minimalista": {
        "primaria": RGBColor(20, 20, 20),
        "secundaria": RGBColor(80, 80, 80),
        "texto": RGBColor(40, 40, 40),
        "sutil": RGBColor(140, 140, 140),
        "fonte_titulo": "Inter",
        "fonte_corpo": "Inter",
    },
    "executivo": {
        "primaria": RGBColor(15, 15, 15),
        "secundaria": RGBColor(60, 60, 60),
        "texto": RGBColor(25, 25, 25),
        "sutil": RGBColor(120, 120, 120),
        "fonte_titulo": "Georgia",
        "fonte_corpo": "Garamond",
    },
    "criativo": {
        "primaria": RGBColor(0, 120, 180),
        "secundaria": RGBColor(0, 160, 140),
        "texto": RGBColor(25, 25, 35),
        "sutil": RGBColor(110, 130, 150),
        "fonte_titulo": "Montserrat",
        "fonte_corpo": "Open Sans",
    },
}


def set_cell_bg(cell, hex_color: str):
    """Define cor de fundo de uma célula."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tcPr.append(shd)


def add_horizontal_rule(doc, cor: RGBColor):
    """Adiciona linha horizontal abaixo do parágrafo atual."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(4)
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "4")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), f"{cor.red:02X}{cor.green:02X}{cor.blue:02X}")
    pBdr.append(bottom)
    pPr.append(pBdr)
    return p


def add_section_title(doc, titulo: str, paleta: dict):
    """Adiciona título de seção com linha inferior."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(14)
    p.paragraph_format.space_after = Pt(0)
    run = p.add_run(titulo.upper())
    run.bold = True
    run.font.name = paleta["fonte_corpo"]
    run.font.size = Pt(10)
    run.font.color.rgb = paleta["primaria"]
    add_horizontal_rule(doc, paleta["primaria"])


def build_cabecalho(doc, dados: dict, paleta: dict):
    """Monta o cabeçalho do currículo."""
    nome = doc.add_paragraph()
    nome.paragraph_format.space_before = Pt(0)
    nome.paragraph_format.space_after = Pt(2)
    run = nome.add_run(dados.get("nome", "").upper())
    run.bold = True
    run.font.name = paleta["fonte_titulo"]
    run.font.size = Pt(24)
    run.font.color.rgb = paleta["primaria"]

    # Linha de contatos
    contatos = []
    if dados.get("telefone"):
        contatos.append(dados["telefone"])
    if dados.get("email"):
        contatos.append(dados["email"])
    if dados.get("cidade"):
        contatos.append(dados["cidade"])
    if dados.get("linkedin"):
        contatos.append(dados["linkedin"])
    if dados.get("portfolio"):
        contatos.append(dados["portfolio"])

    if contatos:
        p = doc.add_paragraph(" · ".join(contatos))
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(4)
        for run in p.runs:
            run.font.name = paleta["fonte_corpo"]
            run.font.size = Pt(9)
            run.font.color.rgb = paleta["sutil"]


def build_resumo(doc, resumo: str, paleta: dict):
    add_section_title(doc, "Resumo Profissional", paleta)
    p = doc.add_paragraph(resumo)
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(4)
    for run in p.runs:
        run.font.name = paleta["fonte_corpo"]
        run.font.size = Pt(10)
        run.font.color.rgb = paleta["texto"]


def build_competencias(doc, competencias: list, paleta: dict):
    add_section_title(doc, "Competências-Chave", paleta)
    # Grid de competências em linha separada por | 
    linha = "  ·  ".join(competencias)
    p = doc.add_paragraph(linha)
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(4)
    for run in p.runs:
        run.font.name = paleta["fonte_corpo"]
        run.font.size = Pt(9.5)
        run.font.color.rgb = paleta["secundaria"]


def build_experiencias(doc, experiencias: list, paleta: dict):
    add_section_title(doc, "Experiência Profissional", paleta)
    for exp in experiencias:
        # Cargo | Empresa | Período
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(8)
        p.paragraph_format.space_after = Pt(1)
        cargo_run = p.add_run(exp.get("cargo", ""))
        cargo_run.bold = True
        cargo_run.font.name = paleta["fonte_corpo"]
        cargo_run.font.size = Pt(10.5)
        cargo_run.font.color.rgb = paleta["primaria"]

        empresa_run = p.add_run(f"  —  {exp.get('empresa', '')}")
        empresa_run.italic = True
        empresa_run.font.name = paleta["fonte_corpo"]
        empresa_run.font.size = Pt(10)
        empresa_run.font.color.rgb = paleta["secundaria"]

        periodo_run = p.add_run(f"  |  {exp.get('periodo', '')}")
        periodo_run.font.name = paleta["fonte_corpo"]
        periodo_run.font.size = Pt(9)
        periodo_run.font.color.rgb = paleta["sutil"]

        # Bullets de atividades
        for bullet in exp.get("atividades", []):
            pb = doc.add_paragraph(style="List Bullet")
            pb.paragraph_format.space_before = Pt(1)
            pb.paragraph_format.space_after = Pt(1)
            pb.paragraph_format.left_indent = Cm(0.4)
            run_b = pb.add_run(bullet)
            run_b.font.name = paleta["fonte_corpo"]
            run_b.font.size = Pt(9.5)
            run_b.font.color.rgb = paleta["texto"]


def build_formacao(doc, formacoes: list, paleta: dict):
    add_section_title(doc, "Formação Acadêmica", paleta)
    for form in formacoes:
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(6)
        p.paragraph_format.space_after = Pt(1)
        curso_run = p.add_run(form.get("curso", ""))
        curso_run.bold = True
        curso_run.font.name = paleta["fonte_corpo"]
        curso_run.font.size = Pt(10)
        curso_run.font.color.rgb = paleta["primaria"]

        inst_run = p.add_run(f"  —  {form.get('instituicao', '')}  |  {form.get('ano', '')}")
        inst_run.font.name = paleta["fonte_corpo"]
        inst_run.font.size = Pt(9.5)
        inst_run.font.color.rgb = paleta["sutil"]


def build_certificacoes(doc, certs: list, paleta: dict):
    if not certs:
        return
    add_section_title(doc, "Certificações", paleta)
    for cert in certs:
        p = doc.add_paragraph(style="List Bullet")
        p.paragraph_format.space_before = Pt(2)
        run = p.add_run(f"{cert.get('nome', '')} — {cert.get('instituicao', '')} ({cert.get('ano', '')})")
        run.font.name = paleta["fonte_corpo"]
        run.font.size = Pt(9.5)
        run.font.color.rgb = paleta["texto"]


def build_idiomas(doc, idiomas: list, paleta: dict):
    if not idiomas:
        return
    add_section_title(doc, "Idiomas", paleta)
    linha = "   ·   ".join(idiomas)
    p = doc.add_paragraph(linha)
    p.paragraph_format.space_before = Pt(6)
    for run in p.runs:
        run.font.name = paleta["fonte_corpo"]
        run.font.size = Pt(9.5)
        run.font.color.rgb = paleta["texto"]


def gerar_curriculo(dados: dict, saida: str, estilo: str = "corporativo"):
    paleta = ESTILOS.get(estilo, ESTILOS["corporativo"])

    doc = Document()

    # Margens do documento
    section = doc.sections[0]
    section.page_height = Cm(29.7)
    section.page_width = Cm(21.0)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.0)
    section.top_margin = Cm(2.0)
    section.bottom_margin = Cm(2.0)

    # Construção das seções
    build_cabecalho(doc, dados, paleta)
    
    if dados.get("resumo"):
        build_resumo(doc, dados["resumo"], paleta)
    
    if dados.get("competencias"):
        build_competencias(doc, dados["competencias"], paleta)
    
    if dados.get("experiencias"):
        build_experiencias(doc, dados["experiencias"], paleta)
    
    if dados.get("formacao"):
        build_formacao(doc, dados["formacao"], paleta)
    
    if dados.get("certificacoes"):
        build_certificacoes(doc, dados["certificacoes"], paleta)
    
    if dados.get("idiomas"):
        build_idiomas(doc, dados["idiomas"], paleta)

    doc.save(saida)
    print(f"✅ Currículo gerado com sucesso: {saida}")
    print(f"   Estilo: {estilo} | Seções: {len([s for s in ['resumo','competencias','experiencias','formacao','certificacoes','idiomas'] if dados.get(s)])}")


def main():
    parser = argparse.ArgumentParser(description="Gerador de Currículo .docx")
    parser.add_argument("--dados", required=True, help="Caminho para o arquivo JSON com os dados")
    parser.add_argument("--saida", default="curriculo.docx", help="Arquivo de saída .docx")
    parser.add_argument("--estilo", default="corporativo",
                        choices=["minimalista", "corporativo", "executivo", "criativo"],
                        help="Estilo visual do currículo")
    args = parser.parse_args()

    dados_path = Path(args.dados)
    if not dados_path.exists():
        print(f"❌ Arquivo de dados não encontrado: {args.dados}")
        sys.exit(1)

    with open(dados_path, "r", encoding="utf-8") as f:
        dados = json.load(f)

    gerar_curriculo(dados, args.saida, args.estilo)


if __name__ == "__main__":
    main()
